"""Document generation utilities for DOCX and PDF outputs."""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Dict, Optional

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generate documents from simple templates."""

    def __init__(
        self,
        template_dir: Optional[str] = None,
        output_dir: Optional[str] = None,
        logger_instance: Optional[logging.Logger] = None,
    ) -> None:
        self.logger = logger_instance or logging.getLogger(__name__)
        self.template_dir = Path(
            template_dir
            or os.getenv("MIA_DOC_TEMPLATE_DIR", "templates/documents")
        )
        self.output_dir = Path(
            output_dir or os.getenv("MIA_DOC_OUTPUT_DIR", "output/documents")
        )
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger.debug(
            "DocumentGenerator initialized template=%s output=%s",
            self.template_dir,
            self.output_dir,
        )

    def create_docx(
        self,
        template_name: str = "proposal",
        context: Optional[Dict[str, str]] = None,
        output_path: Optional[str] = None,
    ) -> str:
        """Create a DOCX document from template data."""

        try:
            from docx import Document  # type: ignore
        except Exception:
            return "python-docx is not installed. Run: pip install python-docx"

        context = context or {}
        output_path = output_path or self._default_filename(
            template_name, "docx"
        )
        document = Document()

        heading = context.get("title", template_name.title())
        document.add_heading(heading, level=0)

        summary = context.get(
            "summary",
            "Documento gerado automaticamente pelo M.I.A. Atualize com detalhes específicos.",
        )
        document.add_paragraph(summary)

        for section_name, section_content in sorted(context.items()):
            if section_name in {"title", "summary"}:
                continue
            document.add_heading(
                section_name.replace("_", " ").title(), level=1
            )
            document.add_paragraph(str(section_content))

        document.save(output_path)
        self.logger.info("DOCX created: %s", output_path)
        return f"DOCX generated at {output_path}"

    def create_pdf(
        self,
        template_name: str = "proposal",
        context: Optional[Dict[str, str]] = None,
        output_path: Optional[str] = None,
    ) -> str:
        """Create a PDF document using reportlab."""

        try:
            from reportlab.lib.pagesizes import LETTER  # type: ignore
            from reportlab.pdfgen import canvas  # type: ignore
        except Exception:
            return "reportlab is not installed. Run: pip install reportlab"

        context = context or {}
        output_path = output_path or self._default_filename(
            template_name, "pdf"
        )
        c = canvas.Canvas(output_path, pagesize=LETTER)

        width, height = LETTER
        x_margin = 72
        y = int(height - 72)

        c.setFont("Helvetica-Bold", 16)
        c.drawString(x_margin, y, context.get("title", template_name.title()))
        y -= 36

        c.setFont("Helvetica", 12)
        summary = context.get(
            "summary",
            "Documento gerado automaticamente pelo M.I.A. Edite conforme necessário.",
        )
        y = self._draw_wrapped_text(c, summary, x_margin, y)

        for section_name, section_content in sorted(context.items()):
            if section_name in {"title", "summary"}:
                continue
            c.setFont("Helvetica-Bold", 13)
            c.drawString(x_margin, y, section_name.replace("_", " ").title())
            y -= 20
            c.setFont("Helvetica", 11)
            y = self._draw_wrapped_text(c, str(section_content), x_margin, y)

        c.showPage()
        c.save()
        self.logger.info("PDF created: %s", output_path)
        return f"PDF generated at {output_path}"

    def _default_filename(self, template_name: str, extension: str) -> str:
        base_name = (
            f"{template_name}_{os.getenv('USER', 'mia')}_{self._timestamp()}"
        )
        return str(self.output_dir / f"{base_name}.{extension}")

    @staticmethod
    def _timestamp() -> str:
        import datetime

        return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    def _draw_wrapped_text(self, canvas, text: str, x: int, y: int) -> int:
        """Render text with simple word wrapping."""

        max_width = 68
        words = text.split()
        line = []
        for word in words:
            line.append(word)
            if len(" ".join(line)) > max_width:
                canvas.drawString(x, y, " ".join(line[:-1]))
                y -= 16
                line = [word]
        if line:
            canvas.drawString(x, y, " ".join(line))
            y -= 16
        return y - 10


__all__ = ["DocumentGenerator"]
